import pandas as pd
from docx import Document
from docx.shared import Pt
import re
import os

def batch_excel_to_word(folder_path, output_file):
    # 1. 获取文件夹下所有的 .xls 文件
    files = [f for f in os.listdir(folder_path) if f.endswith('.xls')]
    
    # 2. 简单的文件名排序逻辑（确保导论在前，数字章节按顺序）
    # 建议：如果文件名是第一章、第二章，程序会按字母排。
    # 我们这里手动指定一个优先级或直接排序
    files.sort() 
    if "导论.xls" in files: # 把导论放到最前面
        files.insert(0, files.pop(files.index("导论.xls")))

    # 创建一个主 Word 文档
    doc = Document()
    doc.add_heading('马克思主义基本原理 复习题库全集', 0)

    # 题目类型顺序
    type_order = ['单选题', '多选题', '判断题', '填空题']

    print(f"检测到 {len(files)} 个文件，开始合并转换...")

    for file_name in files:
        file_path = os.path.join(folder_path, file_name)
        chapter_name = file_name.replace('.xls', '')
        print(f"正在处理: {chapter_name}")

        try:
            # 读取Excel
            df = pd.read_excel(file_path)
            
            # 为每个章节添加一个大标题
            doc.add_heading(chapter_name, level=1)

            for q_type in type_order:
                subset = df[df['题目类型'] == q_type]
                if subset.empty:
                    continue

                # 添加题目类型子标题
                doc.add_heading(q_type, level=2)

                for index, row in subset.iterrows():
                    raw_question = str(row['大题题干'])
                    p = doc.add_paragraph(style='List Number')
                    
                    if q_type == '填空题':
                        # --- 填空题优化逻辑 ---
                        answer_text = str(row['选项A']) if pd.notna(row['选项A']) else "____"
                        # 兼容中英文括号
                        if "()" in raw_question or "（）" in raw_question:
                            parts = re.split(r'\(\)|（）', raw_question)
                            for i, part in enumerate(parts):
                                p.add_run(part)
                                if i < len(parts) - 1:
                                    res_run = p.add_run(f" {answer_text} ")
                                    res_run.underline = True
                                    res_run.bold = True
                        else:
                            p.add_run(raw_question)
                            doc.add_paragraph(f"【答案】：{answer_text}")

                    elif q_type in ['单选题', '多选题']:
                        p.add_run(raw_question)
                        # 循环添加 A-F 选项
                        for char in ['A', 'B', 'C', 'D', 'E', 'F']:
                            col = f'选项{char}'
                            if col in df.columns and pd.notna(row[col]):
                                doc.add_paragraph(f"{char}. {row[col]}", style='List Bullet 2')
                        
                        ans_p = doc.add_paragraph()
                        ans_p.add_run(f"正确答案：{row['正确答案']}").font.size = Pt(10)

                    elif q_type == '判断题':
                        p.add_run(raw_question)
                        # A通常代表正确，B代表错误
                        ans_text = "正确" if str(row['正确答案']).strip().upper() == 'A' else "错误"
                        doc.add_paragraph(f"【判断】：{ans_text}")

                    # 每道题后加一个小分隔
                    # doc.add_paragraph("-" * 10) 

            # 每个章节结束后分页（可选）
            doc.add_page_break()

        except Exception as e:
            print(f"处理文件 {file_name} 时出错: {e}")

    # 保存最终文档
    doc.save(output_file)
    print(f"\n全部转换完成！文件保存为: {output_file}")

if __name__ == "__main__":
    # "." 代表当前运行脚本的文件夹
    current_dir = "." 
    output_filename = "马原复习题库汇总_全章节版.docx"
    
    batch_excel_to_word(current_dir, output_filename)